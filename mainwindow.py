import json
import os
import sys

import cv2
import numpy as np
from PyQt5.QtCore import Qt, QPoint, QTimer, QPointF, QRectF, pyqtSignal
from PyQt5.QtGui import QImage, QPixmap, QTextCharFormat, QFont, QTextCursor, QTransform, QTextImageFormat, \
    QFontDatabase, QTextListFormat, QTextBlockFormat, QColor
from PyQt5.QtPrintSupport import QPrinter, QPrintPreviewDialog
from PyQt5.QtWidgets import QMainWindow, QFileDialog, QMenu, QAction, QColorDialog, QCheckBox, QSpinBox
from docx import Document
from docx2pdf import convert
from ultralytics import YOLO

import common
from common import *
from UI.mainUI import Ui_MainWindow
from det_thread import DetThread
from toolUI.CheckableComboBox import replace_comboBox_with_checkable
from toolUI.ImageLabel import convert_to_imagelabel, BoundingBox
from toolUI.ImageTextEdit import replace_textedit_with_imagetextedit, ImageTextEdit
from toolUI.TipsMessageBox import TipsMessageBox
from toolUI.cameranums import Camera


class MainWindow(QMainWindow, Ui_MainWindow):
    """
    主窗口类，负责初始化和管理UI界面。

    Attributes:
        det_thread (DetThread): 检测线程实例。
    """
    texteditUpdated = pyqtSignal(ImageTextEdit)

    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.setupUi(self)
        self.m_flag = False
        self.textedit_image = {}

        # self.setWindowFlags(Qt.Window | Qt.FramelessWindowHint
        #                     | Qt.WindowSystemMenuHint | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)
        self.setWindowFlags(Qt.FramelessWindowHint)

        self.det_thread = DetThread()
        self.connect_signals()
        self.detect_area_initUI()
        self.report_area_initUI()
        self.repaint_area_initUI()
        self.load_setting()
        self.preheat_models()

    def detect_area_initUI(self):
        """
        推理区信号槽链接
        """
        self.maxButton.setCheckable(True)
        self.maxButton.clicked.connect(self.max_or_restore)

        self.comboBox_model = replace_comboBox_with_checkable(self.comboBox_model)
        self.comboBox_task.addItems(["骨龄评估", "检测", "分类"])
        self.models = {}
        self.model_list = []
        self.update_model_list()
        self.comboBox_task.currentTextChanged.connect(self.update_model_items)
        self.comboBox_model.currentTextChanged.connect(self.select_model)

        self.fileButton.clicked.connect(self.open_file)
        self.cameraButton.clicked.connect(self.chose_cam)
        self.runButton.clicked.connect(self.run_or_continue)
        self.runButton.setChecked(False)
        self.detectButton.clicked.connect(self.detect)

        self.confSpinBox.valueChanged.connect(lambda x: self.change_val(x, 'confSpinBox'))
        self.confSlider.valueChanged.connect(lambda x: self.change_val(x, 'confSlider'))
        self.iouSpinBox.valueChanged.connect(lambda x: self.change_val(x, 'iouSpinBox'))
        self.iouSlider.valueChanged.connect(lambda x: self.change_val(x, 'iouSlider'))
        self.rateSpinBox.valueChanged.connect(lambda x: self.change_val(x, 'rateSpinBox'))
        self.rateSlider.valueChanged.connect(lambda x: self.change_val(x, 'rateSlider'))

        self.qtimer = QTimer(self)
        self.qtimer.setSingleShot(True)
        self.qtimer.timeout.connect(lambda: self.label_bottom.clear())

        self.checkBox_enable.clicked.connect(self.checkrate)
        self.checkBox_male.toggled.connect(self.gender_checked)
        self.checkBox_female.toggled.connect(self.gender_checked)

        self.ProgressSlider.sliderPressed.connect(self.on_slider_pressed)
        self.ProgressSlider.sliderReleased.connect(self.on_slider_released)
        self.ProgressSlider.sliderMoved.connect(self.on_slider_moved)

    def report_area_initUI(self):
        """
        报告区信号槽链接
        """

        self.newButton.clicked.connect(self.newFile)
        self.openButton.clicked.connect(self.openFile)
        self.saveButton.clicked.connect(self.saveFile)
        self.saveasButton.clicked.connect(self.saveAsFile)
        self.printButton.clicked.connect(self.printPreview)

        fontFamilies = QFontDatabase().families()
        fontFamilies.sort()
        self.comboBox_font.addItems(fontFamilies)
        self.comboBox_font.currentTextChanged.connect(self.fontFamilyChanged)
        self.comboBox_size.addItems([str(i) for i in range(6, 72, 2)])
        self.comboBox_size.setCurrentText("12")
        self.comboBox_size.currentIndexChanged.connect(self.fontSizeChanged)
        self.comboBox_size.lineEdit().editingFinished.connect(self.fontSizeChanged)
        self.colorButton.clicked.connect(self.selectColor)

        self.boldButton.clicked.connect(self.toggleBold)
        self.italicButton.clicked.connect(self.toggleItalic)
        self.underlineButton.clicked.connect(self.toggleUnderline)

        self.alignleftButton.clicked.connect(lambda: self.textEdit.setAlignment(Qt.AlignLeft))
        self.aligncenterButton.clicked.connect(lambda: self.textEdit.setAlignment(Qt.AlignCenter))
        self.alignrightButton.clicked.connect(lambda: self.textEdit.setAlignment(Qt.AlignRight))
        self.alignjustifyButton.clicked.connect(lambda: self.textEdit.setAlignment(Qt.AlignJustify))

        self.bulletlistButton.clicked.connect(self.toggleBulletList)
        self.numberedlistButton.clicked.connect(self.toggleNumberedList)
        self.imageButton.clicked.connect(self.insertImage)
        self.angle_count = 0
        self.rotateleftButton.clicked.connect(lambda: self.process_Image(-1))
        self.rotaterightButton.clicked.connect(lambda: self.process_Image(1))
        self.imageZoomInButton.clicked.connect(lambda: self.process_Image(1.1))
        self.imageZoomOutButton.clicked.connect(lambda: self.process_Image(0.9))

        self.textEdit = replace_textedit_with_imagetextedit(self.textEdit)
        self.texteditUpdated.connect(self.update_textedit_reference)

    def repaint_area_initUI(self):
        """
        重绘区信号槽链接
        """
        self.repaint_label = convert_to_imagelabel(self.repaint_label)
        self.repaint_label.setFocusPolicy(Qt.StrongFocus)
        self.repaint_label.mousePressEvent = self.on_mouse_press
        self.repaint_label.mouseMoveEvent = self.on_mouse_move
        self.repaint_label.mouseReleaseEvent = self.on_mouse_release

        # 管理多个标注框
        self.current_box = None
        self.dragging_box = None
        self.dragging_offset = None

        self.tabWidget_repaint.currentChanged.connect(self.on_tab_change)
        self.connect_checkbox_spinbox()
        self.openimageButton.clicked.connect(self.open_image)
        self.openimageButton_2.clicked.connect(self.open_image)
        self.clearButton.clicked.connect(self.clear_current_category_boxes)
        self.clearButton_2.clicked.connect(self.clear_current_category_boxes)
        self.clearallButton.clicked.connect(self.clear_all_boxes)
        self.clearallButton_2.clicked.connect(self.clear_all_boxes)
        self.saveclassifyButton.clicked.connect(self.save_cropped_images)
        self.savedetectButton.clicked.connect(self.save_annotations)
        self.reportButton.clicked.connect(self.re_report)

    # detect_events
    def connect_signals(self):
        """
        连接子线程。
        """
        self.det_thread.update_label_current.connect(lambda img: self.show_image(img, self.label_current))
        self.det_thread.update_label_previous.connect(lambda img: self.show_image(img, self.label_previous))
        self.det_thread.progress_updated.connect(self.update_progress)
        self.det_thread.update_fps.connect(self.update_fps_label)
        self.det_thread.update_bottom_info.connect(self.bottom_msg)
        self.det_thread.update_report_info.connect(self.report_info.setPlainText)
        self.det_thread.update_detection_info.connect(self.detection_info.setPlainText)

    def preheat_models(self):
        """
        预热模型。
        """
        for model_file in self.model_list:
            model_path = os.path.join('./model', model_file)
            self.models[model_file] = YOLO(model_path)

    def update_model_list(self):
        """
        更新模型列表。
        """
        model_list = os.listdir('./model')
        model_list = [file for file in model_list if file.endswith('.pt')]
        model_list.sort(key=lambda x: os.path.getsize('./model/' + x))

        if model_list != self.model_list:
            self.model_list = model_list
            self.update_model_items()

    def update_model_items(self):
        """
        更新模型项。
        """
        selected_task = self.comboBox_task.currentText()
        self.comboBox_model.clear()

        if selected_task == "骨龄评估":
            cls_models = [file for file in self.model_list if file.endswith('-cls.pt')]
            det_models = [file for file in self.model_list if file.endswith('-det.pt')]
            self.comboBox_model.setMultiSelect(True)
            self.comboBox_model.addItems(cls_models + det_models)
        elif selected_task == "检测":
            filtered_list = [file for file in self.model_list if file.endswith('-det.pt')]
            self.comboBox_model.setMultiSelect(False)
            self.comboBox_model.addItems(filtered_list)
        elif selected_task == "分类":
            filtered_list = [file for file in self.model_list if file.endswith('-cls.pt')]
            self.comboBox_model.setMultiSelect(False)
            self.comboBox_model.addItems(filtered_list)

    def select_model(self, model_name):
        """
        选择模型。
        """
        self.bottom_msg(f'Select model to {model_name}')

    def open_file(self):
        """
        打开文件对话框以选择视频或图像文件。
        """
        config_file = 'config/file_default_path.json'
        config = json.load(open(config_file, 'r', encoding='utf-8'))
        open_fold = config['detect_default_path']
        if not os.path.exists(open_fold):
            open_fold = os.getcwd()
        name, _ = QFileDialog.getOpenFileName(self, 'Video/image', open_fold,
                                              "video/image file(*.mp4 *.mkv *.avi *.flv *.jpg *.png)")
        try:
            if name:
                self.det_thread.stop()
                self.det_thread.wait()
                self.label_previous.clear()
                self.label_current.clear()
                self.det_thread.source = name
                self.ProgressSlider.setValue(0)
                self.det_thread.start()
                self.runButton.setChecked(True)
                self.bottom_msg(f'Loaded file：{os.path.basename(name)}')
                config['detect_default_path'] = os.path.dirname(name)
                config_json = json.dumps(config, ensure_ascii=False, indent=2)
                with open(config_file, 'w', encoding='utf-8') as f:
                    f.write(config_json)
                if name.lower().endswith(('.jpg', '.png')):
                    self.open_image(name)
        except Exception as e:
            self.bottom_msg(f'{e}')

    def chose_cam(self):
        """
        选择相机源。
        """
        if self.cameraButton.isChecked():
            try:
                self.show_tips('Loading camera')
                _, cams = Camera().get_cam_num()
                popMenu = QMenu()
                popMenu.setFixedWidth(self.cameraButton.width())
                popMenu.setStyleSheet('''
                                        QMenu {
                                        font-size: 16px;
                                        font-family: "Microsoft YaHei UI";
                                        font-weight: light;
                                        color:white;
                                        border-style: solid;
                                        border-width: 0px;
                                        border-color: rgba(255, 255, 255, 255);
                                        border-radius: 3px;
                                        background-color: rgba(200, 200, 200,50);}
                                    ''')

                for cam in cams:
                    action = QAction(str(cam), self)
                    popMenu.addAction(action)

                action = popMenu.exec_(self.calculate_menu_position())
                if action:
                    self.det_thread.stop()
                    self.det_thread.wait()
                    self.label_previous.clear()
                    self.label_current.clear()
                    self.det_thread.source = action.text()
                    self.ProgressSlider.setValue(0)
                    self.det_thread.start()
                    self.runButton.setChecked(True)
                    self.bottom_msg(f'Loading camera：{action.text()}')
            except Exception as e:
                self.bottom_msg(f'{e}')
        else:
            self.close_camera()

    def calculate_menu_position(self):
        """
        计算菜单位置。
        """
        x = self.groupBox_input.mapToGlobal(self.cameraButton.pos()).x()
        y = self.groupBox_input.mapToGlobal(self.cameraButton.pos()).y()
        y += self.cameraButton.frameGeometry().height()
        return QPoint(x, y)

    def close_camera(self):
        """
        关闭相机。
        """
        self.det_thread.stop()
        self.det_thread.wait()
        self.label_previous.clear()
        self.label_current.clear()
        self.detection_info.clear()
        self.report_info.clear()
        self.show_tips('Camera closing')

    def change_val(self, x, flag):
        """
        改变参数值。
        """
        if flag == 'confSpinBox':
            self.confSlider.setValue(int(x * 100))
        elif flag == 'confSlider':
            self.confSpinBox.setValue(x / 100)
            self.det_thread.conf = x / 100
        elif flag == 'iouSpinBox':
            self.iouSlider.setValue(int(x * 100))
        elif flag == 'iouSlider':
            self.iouSpinBox.setValue(x / 100)
            self.det_thread.iou = x / 100
        elif flag == 'rateSpinBox':
            self.rateSlider.setValue(x)
        elif flag == 'rateSlider':
            self.rateSpinBox.setValue(x)
            if self.checkBox_enable.isChecked():
                self.det_thread.rate = x
        else:
            pass

    def update_progress(self, frame_number):
        """
        更新进度条。
        """
        if self.det_thread.total_frames > 0:
            self.ProgressSlider.setValue(int((frame_number / self.det_thread.total_frames) * 100))

    def on_slider_pressed(self):
        """
        滑块按下事件处理。
        """
        self.det_thread.pause()
        self.runButton.setChecked(False)

    def on_slider_released(self):
        """
        滑块释放事件处理。
        """
        self.runButton.setChecked(True)
        self.det_thread.resume()

    def on_slider_moved(self, position):
        """
        滑块移动事件处理。
        """
        frame_number = int((position / 100) * self.det_thread.total_frames)
        self.det_thread.set_position(frame_number)

    def run_or_continue(self):
        """
        开始或继续运行。
        """
        if self.runButton.isChecked():
            self.det_thread.resume()
        else:
            self.det_thread.pause()

    def detect(self):
        """
        触发检测操作。
        """
        self.detection_info.clear()
        self.report_info.clear()
        common.classify_info.clear()
        common.detection_info.clear()

        self.det_thread.detect_button_status = self.detectButton.isChecked()
        selected_models = self.comboBox_model.selected_items

        if not selected_models:
            selected_task = self.comboBox_task.currentText()
            selected_models = self.get_default_models(selected_task)

            self.comboBox_model.selected_items = selected_models

        self.det_thread.model_paths = selected_models
        self.det_thread.models = {path: self.models[path] for path in selected_models}
        self.det_thread.is_male = self.checkBox_male.isChecked()
        self.det_thread.task = self.comboBox_task.currentText()

        if not self.det_thread.isRunning():
            self.det_thread.start()

    def get_default_models(self, selected_task):
        """
        获取默认模型。
        """
        if selected_task == "骨龄评估":
            cls_models = [file for file in self.model_list if file.endswith('-cls.pt')]
            det_models = [file for file in self.model_list if file.endswith('-det.pt')]
            return [cls_models[0], det_models[0]] if cls_models and det_models else []
        filtered_list = [file for file in self.model_list if '-n' in file]
        return [filtered_list[0]] if filtered_list else []

    def checkrate(self):
        """
        检查并设置播放速率。
        """
        self.det_thread.rate = self.rateSpinBox.value() if self.checkBox_enable.isChecked() else 0

    def gender_checked(self):
        """
        性别选择切换处理。
        """
        if self.sender() == self.checkBox_male and self.checkBox_male.isChecked():
            self.checkBox_female.setChecked(False)
        elif self.sender() == self.checkBox_female and self.checkBox_female.isChecked():
            self.checkBox_male.setChecked(False)
        self.save_setting()

    def update_fps_label(self, fps):
        """
        更新FPS标签。
        """
        self.fps_label.setText(f"FPS: {fps:.2f}")

    def stop(self):
        """
        停止检测线程。
        """
        self.det_thread.stop()
        self.det_thread.wait()
        self.label_previous.clear()
        self.checkBox_autosave.setEnabled(True)

    @staticmethod
    def show_image(img_src, label):
        """
        显示图像在指定标签上。

        Args:
            img_src (np.ndarray): 图像数据。
            label (QLabel): 用于显示图像的标签。
        """
        try:
            ih, iw, _ = img_src.shape
            w = label.width()
            h = label.height()

            if iw / w > ih / h:
                scal = w / iw
                nw = w
                nh = int(scal * ih)
            else:
                scal = h / ih
                nw = int(scal * iw)
                nh = h

            img_src_ = cv2.resize(img_src, (nw, nh))
            frame = cv2.cvtColor(img_src_, cv2.COLOR_BGR2RGB)
            img = QImage(frame.data, frame.shape[1], frame.shape[0], frame.shape[1] * frame.shape[2],
                         QImage.Format_RGB888)
            pixmap = QPixmap.fromImage(img)

            label.setAlignment(Qt.AlignCenter)
            label.setPixmap(pixmap)
        except Exception as e:
            print(repr(e))

    # repaint_events
    def on_tab_change(self):
        """处理选项卡切换事件"""
        # 清除所有方框
        BoundingBox.bounding_boxes.clear()
        self.repaint_label.update()

    def connect_checkbox_spinbox(self):
        self.checkboxes_9 = []
        self.checkboxes_7 = []
        checkboxes_9 = self.groupBox_repaint_cls.findChildren(QCheckBox)
        spinboxes_9 = self.groupBox_repaint_cls.findChildren(QSpinBox)
        checkboxes_7 = self.groupBox_repaint_det.findChildren(QCheckBox)
        for i, (color, name) in enumerate(zip(COLORS, CLASSIFY), 1):
            checkboxes_9[i - 1].clicked.connect(self.on_checkbox_checked)
            spinboxes_9[i - 1].valueChanged.connect(self.on_spinbox_change)
            self.checkboxes_9.append((checkboxes_9[i - 1], spinboxes_9[i - 1], color, name))
        for i, (color, name) in enumerate(zip(COLORS, DETECT), 1):
            checkboxes_7[i - 1].clicked.connect(self.on_checkbox_checked)
            self.checkboxes_7.append((checkboxes_7[i - 1], color, name))

    def on_checkbox_checked(self):
        if self.tabWidget_repaint.currentIndex() == 0:  # 9个类别的选项卡
            for checkbox, spinbox, color, name in self.checkboxes_9:
                if checkbox.isChecked() and checkbox != self.sender():
                    checkbox.setChecked(False)
                elif checkbox.isChecked() and checkbox == self.sender():
                    self.current_category_color = QColor(*color)  # 设置当前类别颜色
                    self.current_category_name = f"{name}_{spinbox.value()}"  # 设置当前类别名称
        else:  # 7个类别的选项卡
            for checkbox, color, name in self.checkboxes_7:
                if checkbox.isChecked() and checkbox != self.sender():
                    checkbox.setChecked(False)
                elif checkbox.isChecked() and checkbox == self.sender():
                    self.current_category_color = QColor(*color)  # 设置当前类别颜色
                    self.current_category_name = name  # 设置当前类别名称

    def on_spinbox_change(self):
        spinbox = self.sender()  # 获取触发信号的 QSpinBox 实例
        if BoundingBox.selected_box and spinbox:
            for checkbox, spinbox_item, color, name in self.checkboxes_9:
                if spinbox_item == spinbox:
                    self.current_category_name = f"{name}_{spinbox.value()}"
                    BoundingBox.selected_box.category_name = self.current_category_name
                    break

    def open_image(self, filename):
        """打开图像文件并加载到标签中。"""
        if filename:
            self.repaint_label.setPixmap(QPixmap(filename))
            self.current_image_path = filename
        else:
            file_name, _ = QFileDialog.getOpenFileName(self, "打开图像文件", "", "Image Files (*.png *.jpg)")
            self.repaint_label.setPixmap(QPixmap(file_name))
            self.current_image_path = file_name

    def clear_current_category_boxes(self):
        """清除当前选择类别的所有方框。"""
        BoundingBox.bounding_boxes = [box for box in BoundingBox.bounding_boxes if
                                      box.category_color != self.current_category_color]
        self.repaint_label.update()

    def clear_all_boxes(self):
        """清除所有方框。"""
        BoundingBox.bounding_boxes.clear()
        self.repaint_label.update()

    def save_annotations(self):
        """保存标注信息到文件。"""
        config_file = 'config/file_default_path.json'

        # 读取配置文件
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
        except FileNotFoundError:
            config = {'repaint_default_path': os.getcwd()}

        open_fold = config.get('repaint_default_path', os.getcwd())
        if not os.path.exists(open_fold):
            open_fold = os.getcwd()

        selected_path = QFileDialog.getExistingDirectory(self, "选择保存路径", open_fold)

        # 检查用户是否选择了路径
        if not selected_path:
            return
        try:
            # 检查detect_datasets目录
            detect_dir = os.path.join(selected_path, "detect_datasets")
            os.makedirs(detect_dir, exist_ok=True)

            # 检查images和labels目录
            images_dir = os.path.join(detect_dir, "images")
            labels_dir = os.path.join(detect_dir, "labels")
            os.makedirs(images_dir, exist_ok=True)
            os.makedirs(labels_dir, exist_ok=True)

            # 保存图像到images目录
            image_name = os.path.basename(self.current_image_path)
            image_save_path = os.path.join(images_dir, image_name)
            self.repaint_label.pixmap.save(image_save_path)

            # 保存注释到labels目录
            label_save_path = os.path.join(labels_dir, os.path.splitext(image_name)[0] + ".txt")
            scale_factor = QPointF(self.repaint_label.image_scaled_size.width() / self.repaint_label.pixmap.width(),
                                   self.repaint_label.image_scaled_size.height() / self.repaint_label.pixmap.height())

            with open(label_save_path, "w") as file:
                for box in BoundingBox.bounding_boxes:
                    start_point, end_point = box.to_image_coordinates(scale_factor)
                    x_center = (start_point.x() + end_point.x()) / 2 / self.repaint_label.pixmap.width()
                    y_center = (start_point.y() + end_point.y()) / 2 / self.repaint_label.pixmap.height()
                    width = abs(end_point.x() - start_point.x()) / self.repaint_label.pixmap.width()
                    height = abs(end_point.y() - start_point.y()) / self.repaint_label.pixmap.height()
                    category_index = DETECT.index(box.category_name)
                    file.write(f"{category_index} {x_center} {y_center} {width} {height}\n")

            # 生成data.yaml
            data_yaml_path = os.path.join(detect_dir, "data.yaml")
            with open(data_yaml_path, "w", encoding='utf-8') as yaml_file:
                yaml_file.write(f"nc: {len(DETECT)}\n")
                yaml_file.write(f"names: {DETECT}\n")

            # 更新配置文件中的路径
            config['repaint_default_path'] = selected_path
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)

        except Exception as e:
            self.bottom_msg(f"Error:Could not save file: {e}")

    def save_cropped_images(self):
        """保存每个类别的裁剪图像。"""
        config_file = 'config/file_default_path.json'

        # 读取配置文件
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
        except FileNotFoundError:
            config = {'repaint_default_path': os.getcwd()}

        open_fold = config.get('repaint_default_path', os.getcwd())
        if not os.path.exists(open_fold):
            open_fold = os.getcwd()

        selected_path = QFileDialog.getExistingDirectory(self, "选择保存路径", open_fold)

        # 检查用户是否选择了路径
        if not selected_path:
            return

        try:
            # 检查classify_datasets目录
            classify_dir = os.path.join(selected_path, "classify_datasets")
            os.makedirs(classify_dir, exist_ok=True)

            # 遍历所有BoundingBox，按类别和等级保存裁剪图像
            for box in BoundingBox.bounding_boxes:
                # 确保 category_name 格式正确
                if '_' not in box.category_name:
                    continue

                # 获取类别和等级
                class_level = box.category_name.split('_', 1)

                # 遍历所有9个类别的checkbox和spinbox
                for checkbox, spinbox, color, name in self.checkboxes_9:
                    if class_level[0] == name:
                        category_dir = os.path.join(classify_dir, name)
                        os.makedirs(category_dir, exist_ok=True)

                        # 检查SpinBox值目录
                        value_dir = os.path.join(category_dir, str(class_level[1]))
                        os.makedirs(value_dir, exist_ok=True)

                        # 保存裁剪图像
                        rect = QRectF(box.start_point, box.end_point).normalized()
                        cropped_pixmap = self.repaint_label.pixmap.copy(rect.toRect())

                        # 确保文件名唯一
                        image_count = len(os.listdir(value_dir))
                        cropped_image_path = os.path.join(value_dir, f"{image_count + 1}.png")
                        cropped_pixmap.save(cropped_image_path)

            # 更新配置文件中的路径并保存
            config['repaint_default_path'] = selected_path
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)

        except Exception as e:
            self.bottom_msg(f"Error: Could not save file: {e}")

    def re_report(self):
        """重新生成报告"""
        pass

    def on_mouse_press(self, event):
        """鼠标按下事件处理。"""
        if not self.repaint_label.is_inside_image(event.pos()):
            return

        pos = self.repaint_label.to_image_coordinates(event.pos())
        scale_factor = QPointF(self.repaint_label.image_scaled_size.width() / self.repaint_label.pixmap.width(),
                               self.repaint_label.image_scaled_size.height() / self.repaint_label.pixmap.height())

        if event.button() == Qt.LeftButton:
            for box in BoundingBox.bounding_boxes:
                if box.selected:
                    corner = box.get_corner(event.pos(), scale_factor, self.repaint_label.image_offset)
                    if corner:
                        box.dragging_corner = corner
                        self.dragging_box = box
                        return

            for box in BoundingBox.bounding_boxes:
                if box.contains(event.pos(), scale_factor, self.repaint_label.image_offset):
                    BoundingBox.deselect_all()
                    self.dragging_box = box
                    box.selected = True
                    self.dragging_offset = pos - box.start_point
                    BoundingBox.selected_box = box
                    if self.tabWidget_repaint.currentIndex() == 0:
                        class_level = self.dragging_box.category_name.split('_', 1)
                        for checkbox, spinbox, color, name in self.checkboxes_9:
                            if class_level[0] == name:
                                spinbox.setValue(int(class_level[1]))
                    self.update_status_bar(box, event.pos())
                    self.repaint_label.update()
                    return

            if self.repaint_label.is_inside_image(event.pos()) and not any(
                    box.contains(event.pos(), scale_factor, self.repaint_label.image_offset) for box in
                    BoundingBox.bounding_boxes):
                BoundingBox.deselect_all()
            self.current_box = BoundingBox(pos, pos, self.current_category_name, self.current_category_color)
            self.update_status_bar(self.current_box, event.pos())

    def on_mouse_move(self, event):
        """鼠标移动事件处理。"""
        if not self.repaint_label.is_inside_image(event.pos()):
            return

        pos = self.repaint_label.to_image_coordinates(event.pos())

        if self.current_box:
            self.current_box.end_point = pos
            self.update_status_bar(self.current_box, event.pos())
            self.repaint_label.update()
        elif self.dragging_box:
            scale_factor = QPointF(self.repaint_label.image_scaled_size.width() / self.repaint_label.pixmap.width(),
                                   self.repaint_label.image_scaled_size.height() / self.repaint_label.pixmap.height())
            if self.dragging_box.dragging_corner:
                self.dragging_box.resize(self.dragging_box.dragging_corner, event.pos(), scale_factor,
                                         self.repaint_label.image_offset)
            else:
                offset = pos - (self.dragging_box.start_point + self.dragging_offset)
                self.dragging_box.move(offset)
                self.dragging_offset = pos - self.dragging_box.start_point
            self.update_status_bar(self.dragging_box, event.pos())
            self.repaint_label.update()

    def on_mouse_release(self, event):
        """鼠标释放事件处理。"""
        if self.current_box:
            if self.current_box.width() < 2 or self.current_box.height() < 2:
                BoundingBox.bounding_boxes.remove(self.current_box)
            self.current_box = None
        elif self.dragging_box:
            self.dragging_box.dragging_corner = None
            self.dragging_box = None
            self.dragging_offset = None
        self.repaint_label.update()

    def keyPressEvent(self, event):
        """键盘按下事件处理。"""
        if event.key() in [Qt.Key_Delete, Qt.Key_Backspace]:
            BoundingBox.remove_selected()
            self.repaint_label.update()

    def update_status_bar(self, box, mouse_pos=None):
        """更新状态栏显示的方框信息及鼠标位置。"""
        start_pos = box.start_point
        end_pos = box.end_point

        # 将坐标值精确到两位小数
        start_x = f"{start_pos.x():.2f}"
        start_y = f"{start_pos.y():.2f}"
        end_x = f"{end_pos.x():.2f}"
        end_y = f"{end_pos.y():.2f}"

        self.start_position_label.setText(f"方框起始位置: ({start_x}, {start_y})")
        self.end_position_label.setText(f"方框终止位置: ({end_x}, {end_y})")
        self.box_size_label.setText(f"方框大小: {box.width():.2f} x {box.height():.2f}")

        if mouse_pos:
            image_coords = self.repaint_label.to_image_coordinates(mouse_pos)
            mouse_x = f"{image_coords.x():.2f}"
            mouse_y = f"{image_coords.y():.2f}"
            self.mouse_position_label.setText(f"鼠标位置: ({mouse_x}, {mouse_y})")

    # report_events
    def newFile(self):
        """
        创建一个新的文档，默认创建 .docx 文件，并弹出对话框让用户选择保存路径。
        """
        # 清空当前文本编辑器内容
        self.textEdit.clear()

        # 弹出对话框让用户选择保存路径
        config_file = 'config/file_default_path.json'
        fname, _ = QFileDialog.getSaveFileName(self, 'Create file as', '',
                                               'Word documents (*.docx *.doc);;All files (*)')
        if fname:
            if not fname.endswith('.docx') and not fname.endswith('.doc'):
                fname = f"{fname}.docx"  # 默认创建 .docx 文件

            try:
                # 创建并保存 .docx 文件
                doc = Document()
                doc._body.clear_content()  # 清除原有内容
                for block in self.textEdit.toPlainText().split('\n'):
                    doc.add_paragraph(block)
                doc.save(fname)

                # 记录文件路径
                self.docfilename = fname

                # 更新配置文件中的路径
                try:
                    with open(config_file, 'r', encoding='utf-8') as f:
                        config = json.load(f)
                except (FileNotFoundError, json.JSONDecodeError):
                    config = {}

                config['report_default_path'] = os.path.dirname(fname)

                with open(config_file, 'w', encoding='utf-8') as f:
                    json.dump(config, f, ensure_ascii=False, indent=2)

            except Exception as e:
                self.bottom_msg(f"Error:Could not create file: {e}")

    def openFile(self):
        """
        打开文件对话框以选择 Word、Text 或其他支持的文件，并加载内容。
        """
        config_file = 'config/file_default_path.json'

        # 读取配置文件
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            config = {}

        open_fold = config.get('report_default_path', os.getcwd())
        if not os.path.exists(open_fold):
            open_fold = os.getcwd()

        fname, _ = QFileDialog.getOpenFileName(self, 'Open file', open_fold,
                                               "Word documents (*.docx *.doc);;Text files (*.txt);;All files (*)")
        if fname:
            try:
                if fname.endswith('.txt'):
                    with open(fname, 'r', encoding='utf-8') as file:
                        self.textEdit.setPlainText(file.read())
                else:
                    doc = Document(fname)
                    self.textEdit.clear()
                    for para in doc.paragraphs:
                        if para.text:
                            self.textEdit.append(para.text)

                # 记录文件路径
                self.docfilename = fname

                # 更新配置文件中的路径
                config['report_default_path'] = os.path.dirname(fname)
                with open(config_file, 'w', encoding='utf-8') as f:
                    json.dump(config, f, ensure_ascii=False, indent=2)

            except Exception as e:
                self.bottom_msg(f"Error:Could not open file: {e}")

    def saveFile(self):
        """
        默认保存当前内容为 .docx 文件。
        """
        if hasattr(self, 'docfilename') and self.docfilename:
            current_file = self.docfilename
        else:
            self.saveAsFile()
            return

        try:
            # 清理不必要的临时图片
            self.textEdit.cleanup_temp_images()

            doc = Document()
            doc._body.clear_content()  # 清除原有内容

            # 处理文本和图片
            cursor = self.textEdit.textCursor()
            cursor.movePosition(QTextCursor.Start)

            while True:
                char_format = cursor.charFormat()
                if char_format.isImageFormat():
                    image_format = char_format.toImageFormat()
                    image_path = image_format.name()

                    # 检查路径是否为相对路径并转换为绝对路径
                    if not os.path.isabs(image_path):
                        image_path = os.path.join(os.getcwd(), image_path)  # 转换为绝对路径

                    if os.path.exists(image_path):
                        doc.add_picture(image_path)
                    else:
                        self.bottom_msg(f"Warning: Image not found: {image_path}")
                else:
                    text = cursor.block().text()
                    doc.add_paragraph(text)

                # 移动到下一个块
                if not cursor.movePosition(QTextCursor.NextBlock):
                    break

            doc.save(current_file)

        except Exception as e:
            self.bottom_msg(f"Error: Could not save file: {e}")

    def saveAsFile(self):
        """
        保存当前内容为指定格式的新文件。
        """
        config_file = 'config/file_default_path.json'

        # 打开文件对话框让用户选择新文件路径
        fname, _ = QFileDialog.getSaveFileName(self, 'Save file as', '',
                                               'Word documents (*.doc *.docx);;Text files (*.txt);;PDF files (*.pdf);;All files (*)')

        if fname:  # 如果用户选择了文件路径
            try:
                # 清理不必要的临时图片
                self.textEdit.cleanup_temp_images()

                # 根据文件扩展名选择保存格式
                file_ext = os.path.splitext(fname)[1].lower()
                if file_ext in ['.doc', '.docx']:
                    doc = Document()
                    doc._body.clear_content()  # 清除原有内容

                    cursor = self.textEdit.textCursor()
                    cursor.movePosition(QTextCursor.Start)

                    while True:
                        char_format = cursor.charFormat()
                        if char_format.isImageFormat():
                            image_format = char_format.toImageFormat()
                            image_path = image_format.name()

                            # 检查路径是否为相对路径并转换为绝对路径
                            if not os.path.isabs(image_path):
                                image_path = os.path.join(os.getcwd(), image_path)

                            if os.path.exists(image_path):
                                doc.add_picture(image_path)
                            else:
                                self.bottom_msg(f"Warning: Image not found: {image_path}")
                        else:
                            text = cursor.block().text()
                            doc.add_paragraph(text)

                        # 移动到下一个块
                        if not cursor.movePosition(QTextCursor.NextBlock):
                            break

                    doc.save(fname)
                elif file_ext == '.txt':
                    with open(fname, 'w', encoding='utf-8') as file:
                        file.write(self.textEdit.toPlainText())
                elif file_ext == '.pdf':
                    doc = Document()
                    doc._body.clear_content()  # 清除原有内容

                    cursor = self.textEdit.textCursor()
                    cursor.movePosition(QTextCursor.Start)

                    while True:
                        char_format = cursor.charFormat()
                        if char_format.isImageFormat():
                            image_format = char_format.toImageFormat()
                            image_path = image_format.name()

                            # 检查路径是否为相对路径并转换为绝对路径
                            if not os.path.isabs(image_path):
                                image_path = os.path.join(os.getcwd(), image_path)

                            if os.path.exists(image_path):
                                doc.add_picture(image_path)
                            else:
                                self.bottom_msg(f"Warning: Image not found: {image_path}")
                        else:
                            text = cursor.block().text()
                            doc.add_paragraph(text)

                        # 移动到下一个块
                        if not cursor.movePosition(QTextCursor.NextBlock):
                            break

                    temp_docx = f"{os.path.splitext(fname)[0]}.docx"
                    doc.save(temp_docx)  # 先保存为 docx 文件
                    convert(temp_docx, fname)  # 转换为 pdf 文件
                    os.remove(temp_docx)  # 删除中间文件
                else:
                    self.bottom_msg("Warning: Unsupported file format.")

                # 更新当前文件路径为用户选择的新文件路径
                self.docfilename = fname

                # 更新配置文件中的路径
                try:
                    with open(config_file, 'r', encoding='utf-8') as f:
                        config = json.load(f)
                except (FileNotFoundError, json.JSONDecodeError):
                    config = {}

                config['report_default_path'] = os.path.dirname(fname)

                with open(config_file, 'w', encoding='utf-8') as f:
                    json.dump(config, f, ensure_ascii=False, indent=2)

            except Exception as e:
                self.bottom_msg(f"Error: Could not save as file: {e}")

    def printPreview(self):
        printer = QPrinter(QPrinter.HighResolution)
        previewDialog = QPrintPreviewDialog(printer, self)
        previewDialog.paintRequested.connect(self.textEdit.print_)
        previewDialog.exec_()

    def selectColor(self):
        color = QColorDialog.getColor()
        if color.isValid():
            self.textEdit.setTextColor(color)

    def toggleBold(self):
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Bold if self.boldButton.isChecked() else QFont.Normal)
        self.mergeFormatOnWordOrSelection(fmt)

    def toggleItalic(self):
        fmt = QTextCharFormat()
        fmt.setFontItalic(self.italicButton.isChecked())
        self.mergeFormatOnWordOrSelection(fmt)

    def toggleUnderline(self):
        fmt = QTextCharFormat()
        fmt.setFontUnderline(self.underlineButton.isChecked())
        self.mergeFormatOnWordOrSelection(fmt)

    def mergeFormatOnWordOrSelection(self, fmt):
        cursor = self.textEdit.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        cursor.mergeCharFormat(fmt)
        self.textEdit.mergeCurrentCharFormat(fmt)

    def fontFamilyChanged(self, font):
        fmt = QTextCharFormat()
        fmt.setFontFamily(font)
        self.mergeFormatOnWordOrSelection(fmt)

    def fontSizeChanged(self):
        fontSize = int(self.comboBox_size.currentText())
        fmt = QTextCharFormat()
        fmt.setFontPointSize(fontSize)
        self.mergeFormatOnWordOrSelection(fmt)

    def toggleBulletList(self):
        cursor = self.textEdit.textCursor()
        cursor.beginEditBlock()

        # 检查段落是否已经有无序列表
        if cursor.currentList() and cursor.currentList().format().style() == QTextListFormat.ListDisc:
            # 取消无序列表
            block_format = QTextBlockFormat()
            block_format.setObjectIndex(-1)  # 将列表索引设置为无效值
            cursor.setBlockFormat(block_format)
        else:
            # 插入无序列表，并确保所有选中的段落都应用相同的列表格式
            list_format = QTextListFormat()
            list_format.setStyle(QTextListFormat.ListDisc)
            cursor.createList(list_format)

        cursor.endEditBlock()

    def toggleNumberedList(self):
        cursor = self.textEdit.textCursor()
        cursor.beginEditBlock()

        # 检查段落是否已经有编号列表
        if cursor.currentList() and cursor.currentList().format().style() == QTextListFormat.ListDecimal:
            # 取消编号列表
            block_format = QTextBlockFormat()
            block_format.setObjectIndex(-1)  # 将列表索引设置为无效值
            cursor.setBlockFormat(block_format)
        else:
            # 插入编号列表，并确保所有选中的段落都应用相同的列表格式
            list_format = QTextListFormat()
            list_format.setStyle(QTextListFormat.ListDecimal)
            cursor.createList(list_format)

        cursor.endEditBlock()

    def insertImage(self):
        fname, _ = QFileDialog.getOpenFileName(self, 'Insert Image', '', 'Images (*.png *.jpg *.bmp);;All files (*)')
        if fname:
            try:
                image = QImage(fname)
                for image_name in list(self.textedit_image.keys()):
                    if image_name in fname:
                        del self.textedit_image[image_name]
                self.insertAndResizeImage(image, fname)
            except Exception as e:
                self.bottom_msg(f"Error: Could not insert image: {e}")

    def process_Image(self, value):
        cursor = self.textEdit.textCursor()
        image_format = cursor.charFormat().toImageFormat()

        if image_format.isValid():
            image_name = image_format.name()
            image = QImage(image_name)
            if not image.isNull():
                # 旋转图像
                image_info = self.textedit_image.get(image_name, {'angle': 0, 'size': (image.width(), image.height())})
                angle = image_info['angle']
                width, height = image_info['size']
                width = int(width)
                height = int(height)
                if width <= 0 or height <= 0:
                    width = image.width()
                    height = image.height()
                if value == 1 or value == -1:
                    new_angle = (angle + value * 90) % 360
                    transform = QTransform().rotate(new_angle)
                    process_image = image.scaled(int(width), int(height), Qt.KeepAspectRatio, Qt.SmoothTransformation)
                    process_image = process_image.transformed(transform)
                    self.textedit_image[image_name] = {'angle': new_angle, 'size': (width, height)}
                else:
                    new_width = int(width * value)
                    new_height = int(height * value)
                    transform = QTransform().rotate(angle)
                    process_image = image.scaled(new_width, new_height, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                    process_image = process_image.transformed(transform)
                    self.textedit_image[image_name] = {'angle': angle, 'size': (new_width, new_height)}

                # 如果处理后的图像宽度超过 QTextEdit 的宽度，则进行缩放
                if process_image.width() > self.textEdit.width():
                    process_image = process_image.scaledToWidth(self.textEdit.width())

                # 删除原有图像，并插入处理后的图像
                cursor.deletePreviousChar()
                cursor.insertImage(process_image, image_name)

    def insertAndResizeImage(self, image, image_path=None):
        if image.width() > self.textEdit.width():
            image = image.scaledToWidth(self.textEdit.width())
        cursor = self.textEdit.textCursor()
        image_format = QTextImageFormat()
        image_format.setName(image_path if image_path else "")  # 设置图片路径作为图片标识符
        image_format.setWidth(image.width())
        image_format.setHeight(image.height())
        cursor.insertImage(image_format)

    def update_textedit_reference(self, new_textedit):
        self.textEdit = new_textedit

    # global_events
    def load_setting(self):
        """
        加载设置，仅更新指定键的值，保留其他配置。
        """
        config_file = 'config/setting.json'
        default_config = {
            "iou": 0.20,
            "conf": 0.50,
            "rate": 1,
            "check": 0,
            "male_checked": True,
            "female_checked": False,
            "font-family": "Arial",  # 默认字体
            "font-size": 12  # 默认字体大小
        }

        # 加载现有配置或使用默认配置
        config = self.load_config(config_file, default_config)
        self.apply_config(config)

    def load_config(self, config_file, default_config):
        """
        加载配置文件，如果文件不存在或内容不全，使用默认配置。
        """
        if os.path.exists(config_file):
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
        else:
            config = {}

        # 合并默认配置与现有配置，保留未修改的键值
        updated_config = {**default_config, **config}
        return updated_config

    def write_config(self, config_file, config):
        """
        写入或更新配置文件，仅更新指定的键值。
        """
        current_config = self.load_config(config_file, {})
        current_config.update(config)  # 更新指定键值
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(current_config, f, ensure_ascii=False, indent=2)

    def save_setting(self):
        """
        保存设置，仅更新需要修改的键值。
        """
        config_file = 'config/setting.json'
        config = {
            'iou': self.iouSpinBox.value(),
            'conf': self.confSpinBox.value(),
            'rate': self.rateSpinBox.value(),
            'check': self.checkBox_enable.checkState(),
            'male_checked': self.checkBox_male.isChecked(),
            'female_checked': self.checkBox_female.isChecked(),
            'font-family': self.comboBox_font.currentText(),
            'font-size': self.comboBox_size.currentText()
        }
        self.write_config(config_file, config)

    def apply_config(self, config):
        """
        应用配置，仅更新指定的设置项。
        """
        self.confSpinBox.setValue(config['conf'])
        self.iouSpinBox.setValue(config['iou'])
        self.rateSpinBox.setValue(config['rate'])
        self.checkBox_enable.setCheckState(config['check'])
        self.checkBox_male.setChecked(config['male_checked'])
        self.checkBox_female.setChecked(config['female_checked'])
        self.comboBox_font.setCurrentText(config['font-family'])
        self.comboBox_size.setCurrentText(str(config['font-size']))

    def show_tips(self, text):
        """
        显示提示信息。
        """
        TipsMessageBox(self.closeButton, title='Tips', text=text, time=2000, auto=True).exec_()

    def bottom_msg(self, msg):
        """
        更新底部信息标签。
        """
        self.label_bottom.setText(msg)

    def max_or_restore(self):
        """
        最大化或还原窗口。
        """
        if self.maxButton.isChecked():
            self.window().showMaximized()
        else:
            self.window().showNormal()

    def mousePressEvent(self, event):
        """
        鼠标按下事件处理。
        """
        self.m_Position = event.pos()
        if event.button() == Qt.LeftButton:
            if 0 < self.m_Position.x() < self.groupBox_body.pos().x() + self.groupBox_body.width() and \
                    0 < self.m_Position.y() < self.groupBox_body.pos().y() + self.groupBox_body.height():
                self.m_flag = True

    def mouseMoveEvent(self, QMouseEvent):
        """
        鼠标移动事件处理。
        """
        if Qt.LeftButton and self.m_flag:
            self.move(QMouseEvent.globalPos() - self.m_Position)

    def mouseReleaseEvent(self, QMouseEvent):
        """
        鼠标释放事件处理。
        """
        self.m_flag = False

    def closeEvent(self, event):
        """
        关闭事件处理。
        """
        self.save_setting()
        self.show_tips('Closing the program')
        sys.exit(0)
